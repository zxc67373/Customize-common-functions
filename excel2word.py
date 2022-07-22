from excel_data_process import excel_process
# https://www.justdopython.com/2020/03/19/python-python-docs/
import xlrd
from docx import Document
book = xlrd.open_workbook('OCR首页分类抽取结果.xls')
# 选择页数为第1页
document = Document()
sheet1 = book.sheets()[0]
for j in range(len(sheet1.col_values(1)[:])):
    flag_yuyin = 0
    flag_ocr = 0
    for i in range(4):
        x= sheet1.col_values(i)[j]
        if i == 0 and x:
            document.add_heading(x, level=1)
            continue
        if x  and x[0] == '7' and len('7122386574184140073') == len(x):
            document.add_heading('https://www.douyin.com/video/'+x, level=2)
            flag_yuyin = 1
            continue
        if flag_yuyin == 1:
            document.add_heading('语音转译结果', level=3)
            document.add_paragraph(x)
            flag_ocr =1
            flag_yuyin =0
            continue
        if flag_ocr == 1:
            document.add_heading('OCR识别结果', level=3)
            document.add_paragraph(x)

document.save(r"test.docx")